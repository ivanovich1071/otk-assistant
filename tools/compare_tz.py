"""Сверка требований по изготовлению с эталоном.

Сравнивается то, что должно совпадать точно: состав комплекта, веса, материалы
и строки спецификаций. Формулировки технических требований сравниваются мягко —
их модель переписывает в канонический вид, и дословного совпадения не ждём.

    python tools/compare_tz.py --tz "output/... (ТЗ).docx" --gold "tests/gold/tz/.../tz.docx"
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

import docx

DESIGNATION = re.compile(r"^(?:\d+[.:)]\s*)?([А-ЯA-Z]{2,8}\s+[\d.]{6,}(?:\s+СБ)?)\s*(.*)$")
EMPTY = {"", "-", "–", "—"}
MASS = re.compile(r"Вес\s*[–—-]\s*([^|\n]+)")
MATERIAL = re.compile(r"Материал\s*[–—-]\s*([^|\n]+)")
PER_PIECE = re.compile(r"\s*\(за\s+\d+\s*шт\)")


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _mass(text: str) -> str:
    """Прочерк в графе означает «нет данных» — сравнивать его с пустотой нечестно."""
    value = _clean(PER_PIECE.sub("", text)).rstrip(".")
    return "" if value in EMPTY else value


def read_doc(path: Path) -> dict:
    doc = docx.Document(str(path))
    blocks: dict[str, dict] = {}
    current: str | None = None

    for paragraph in doc.paragraphs:
        text = _clean(paragraph.text)
        if not text:
            continue
        match = DESIGNATION.match(text)
        if match:
            current = _clean(match.group(1))
            blocks.setdefault(current, {"title": _clean(match.group(2)),
                                        "mass": "", "material": "", "lines": []})
            if not blocks[current]["title"]:
                blocks[current]["title"] = _clean(match.group(2))
            continue
        if current is None:
            continue
        block = blocks[current]
        block["lines"].append(text)
        if not block["mass"] and (found := MASS.search(paragraph.text)):
            block["mass"] = _mass(found.group(1))
        if not block["material"] and (found := MATERIAL.search(paragraph.text)):
            block["material"] = _clean(found.group(1))

    rows = set()
    for table in doc.tables:
        for row in table.rows:
            cells = [_clean(c.text) for c in row.cells]
            if len(cells) >= 3 and cells[1] and DESIGNATION.match(cells[1]):
                rows.add((cells[0], cells[1], cells[2],
                          cells[3] if len(cells) > 3 else "",
                          _mass(cells[4]) if len(cells) > 4 else ""))
    return {"blocks": blocks, "rows": rows}


def compare(made: dict, gold: dict) -> list[str]:
    problems: list[str] = []

    missing = sorted(set(gold["blocks"]) - set(made["blocks"]))
    extra = sorted(set(made["blocks"]) - set(gold["blocks"]))
    for key in missing:
        problems.append(f"нет в результате: {key} {gold['blocks'][key]['title']}")
    for key in extra:
        problems.append(f"нет в эталоне: {key} {made['blocks'][key]['title']}")

    for key in sorted(set(made["blocks"]) & set(gold["blocks"])):
        ours, theirs = made["blocks"][key], gold["blocks"][key]
        for field, label in (("mass", "вес"), ("material", "материал")):
            if ours[field] and theirs[field] and ours[field] != theirs[field]:
                problems.append(
                    f"{key}: {label} «{ours[field]}» против «{theirs[field]}» в эталоне")
            elif theirs[field] and not ours[field]:
                problems.append(f"{key}: {label} не заполнен, в эталоне «{theirs[field]}»")

    for row in sorted(gold["rows"] - made["rows"]):
        problems.append("строка спецификации потеряна: " + " | ".join(row))
    return problems


def main() -> None:
    ap = argparse.ArgumentParser(description="Сверка ТЗ с эталоном")
    ap.add_argument("--tz", type=Path, required=True)
    ap.add_argument("--gold", type=Path, required=True)
    ap.add_argument("--quiet", action="store_true", help="только итог")
    args = ap.parse_args()

    made, gold = read_doc(args.tz), read_doc(args.gold)
    problems = compare(made, gold)

    common = len(set(made["blocks"]) & set(gold["blocks"]))
    print(f"позиций в эталоне: {len(gold['blocks'])}, найдено: {common}")
    print(f"строк спецификаций: {len(made['rows'])} против {len(gold['rows'])} в эталоне, "
          f"совпало {len(made['rows'] & gold['rows'])}")
    print(f"расхождений: {len(problems)}")
    if not args.quiet:
        for line in problems:
            print("  ! " + line)


if __name__ == "__main__":
    main()
