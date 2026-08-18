"""Сборка карты обмера (.docx) из markup.json.

Форматирование берётся из templates/karta_obmera.docx — шаблон вырезан из
реальной карты, поэтому шапка, ширины колонок, шрифты и рамки совпадают
с привычными один в один. Строки клонируются из образцов в шаблоне.

Использование:
    python tools/build_docx.py work/markup.json -o "output/... (карта обмера).docx"
"""
from __future__ import annotations

import argparse
import copy
import json
import subprocess
import sys
from pathlib import Path

import docx
from docx.shared import Emu, Pt


TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "karta_obmera.docx"
ROW_HEADER, ROW_GROUP, ROW_DATA = 0, 1, 2
FRAME_HEIGHT_PT = 13     # высота вставляемой рамки допуска ≈ строка текста
TECH_GROUP = "Технические требования чертежа"
# Параметры зубчатого венца обмеряются наравне с диаметрами, поэтому идут
# в карту отдельной секцией, а не выбрасываются вместе со служебными таблицами.
PARAM_GROUP = "Параметры зубчатого венца"


def _clone(table, model_index: int):
    model = table.rows[model_index]._element
    new = copy.deepcopy(model)
    table._tbl.append(new)
    return table.rows[-1]


def _set(cell, text: str) -> None:
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if text:
        p.add_run(text)


def _set_image(cell, image: Path, text: str = "") -> None:
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if text:
        p.add_run(text + " ")
    p.add_run().add_picture(str(image), height=Pt(FRAME_HEIGHT_PT))


def _frame_crop(drawing: Path, item: dict, work: Path) -> Path:
    """Рамку допуска не расшифровываем текстом — вырезаем с чертежа как картинку."""
    work.mkdir(parents=True, exist_ok=True)
    out = work / f"frame_{str(item['no']).replace('-', '_')}.png"
    if not out.exists():
        box = f"{item['x']},{item['y']},{item['w']},{item['h']}"
        subprocess.run(
            [sys.executable, str(Path(__file__).with_name("crop.py")), str(drawing),
             "crops", "-o", str(out), "--box", box, "--scale", "2", "--pad", "3"],
            check=True, capture_output=True,
        )
    return out


def build(markup: dict, out: Path, template: Path = TEMPLATE) -> int:
    doc = docx.Document(str(template))
    table = doc.tables[0]

    title = " ".join(x for x in (markup.get("designation"), markup.get("title")) if x)
    doc.paragraphs[0].add_run(title)

    drawing = Path(markup["drawing"])
    work = out.parent.parent / "work" / "frames"

    rows = 0
    for group in markup["groups"]:
        row = _clone(table, ROW_GROUP)
        _set(row.cells[1], group["name"])
        rows += 1
        for item in group["items"]:
            row = _clone(table, ROW_DATA)
            _set(row.cells[0], str(item.get("no", "")))
            if item.get("kind") == "frame":
                _set_image(row.cells[1], _frame_crop(drawing, item, work),
                           item.get("value", ""))
            else:
                _set(row.cells[1], item.get("value", ""))
            rows += 1

    for section, lines in ((PARAM_GROUP, markup.get("parameters")),
                           (TECH_GROUP, markup.get("tech_requirements"))):
        if not lines:
            continue
        row = _clone(table, ROW_GROUP)
        _set(row.cells[1], section)
        rows += 1
        for line in lines:
            row = _clone(table, ROW_DATA)
            _set(row.cells[0], line.get("no", "") if isinstance(line, dict) else "")
            _set(row.cells[1], line["text"] if isinstance(line, dict) else str(line))
            rows += 1

    # Образцы строк из шаблона больше не нужны.
    for idx in (ROW_DATA, ROW_GROUP):
        el = table.rows[idx]._element
        el.getparent().remove(el)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return rows


def main() -> None:
    ap = argparse.ArgumentParser(description="Карта обмера в Word из markup.json")
    ap.add_argument("markup", type=Path)
    ap.add_argument("-o", "--out", type=Path, required=True)
    ap.add_argument("--template", type=Path, default=TEMPLATE)
    args = ap.parse_args()

    markup = json.loads(args.markup.read_text(encoding="utf-8"))
    n = build(markup, args.out, args.template)
    print(f"строк в таблице: {n} -> {args.out}")


if __name__ == "__main__":
    main()
