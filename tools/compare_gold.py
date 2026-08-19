"""Печать и сверка карт обмера.

    python tools/compare_gold.py --card "output/... .docx"
    python tools/compare_gold.py --card "output/... .docx" --gold "tests/gold/.../card.docx"

Сверка идёт по номеру позиции: что потеряно, что лишнее, где разошлись значения.
Строки-группы сравниваются отдельно — по составу и порядку.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import docx


GROUP_RE = re.compile(
    r"^(Главный вид|Вид\b.*|Разрез\b.*|Сечение\b.*|Выноска\b.*|Вноска\b.*"
    # Заголовок с масштабом идёт в карту как на листе, без придуманного префикса.
    r"|[А-ЯA-Z]\d?\s*\(\s*\d+\s*:\s*[\d,]+\s*\)"
    r"|Параметры\b.*|Технические требования.*)$"
)


def read_card(path: Path) -> tuple[str, list[dict]]:
    doc = docx.Document(str(path))
    title = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
    if not doc.tables:
        raise SystemExit(f"В документе нет таблицы: {path}")
    table = doc.tables[0]

    rows: list[dict] = []
    group = ""
    for i, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if i == 0 or (cells[0] in ("1", "2") and not cells[1] and len(cells) > 2):
            continue  # шапка и подшапка «1 | 2»
        no, value = cells[0], cells[1] if len(cells) > 1 else ""
        has_img = bool(row.cells[1]._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        )) if len(cells) > 1 else False
        if not no and GROUP_RE.match(value):
            group = value
            rows.append({"kind": "group", "no": "", "value": value, "group": value, "img": False})
            continue
        if not no and not value and not has_img:
            continue
        rows.append({"kind": "item", "no": no, "value": value, "group": group, "img": has_img})
    return title, rows


def dump(title: str, rows: list[dict]) -> None:
    print(f"Заголовок: {title}")
    items = [r for r in rows if r["kind"] == "item"]
    groups = [r for r in rows if r["kind"] == "group"]
    print(f"Групп: {len(groups)}, позиций: {len(items)}, "
          f"из них с картинкой рамки: {sum(1 for r in items if r['img'])}")
    for r in rows:
        if r["kind"] == "group":
            print(f"\n[{r['value']}]")
        else:
            mark = "  [рамка]" if r["img"] else ""
            print(f"  {r['no']:>6} | {r['value']}{mark}")


def _norm(s: str) -> str:
    """Пробелы и переносы в карте несущественны — сравниваем смысл, а не вёрстку."""
    return re.sub(r"\s+", "", s)


def compare(card: list[dict], gold: list[dict]) -> int:
    ci = {r["no"]: r for r in card if r["kind"] == "item" and r["no"]}
    gi = {r["no"]: r for r in gold if r["kind"] == "item" and r["no"]}

    missing = [n for n in gi if n not in ci]
    extra = [n for n in ci if n not in gi]
    diff = [(n, gi[n]["value"], ci[n]["value"])
            for n in gi if n in ci and _norm(gi[n]["value"]) != _norm(ci[n]["value"])]
    spacing = sum(1 for n in gi if n in ci
                  and gi[n]["value"] != ci[n]["value"]
                  and _norm(gi[n]["value"]) == _norm(ci[n]["value"]))

    cg = [r["value"] for r in card if r["kind"] == "group"]
    gg = [r["value"] for r in gold if r["kind"] == "group"]

    print(f"\nЭталон: {len(gi)} позиций, карта: {len(ci)} позиций")
    print(f"Совпало номеров: {len(gi) - len(missing)} из {len(gi)}"
          f"  ({100 * (len(gi) - len(missing)) / max(1, len(gi)):.0f}%)")
    same_value = len(gi) - len(missing) - len(diff)
    print(f"Совпало значений: {same_value} из {len(gi)}"
          f"  ({100 * same_value / max(1, len(gi)):.0f}%)"
          f"{f', из них {spacing} — с точностью до пробелов' if spacing else ''}")

    if gg != cg:
        print("\nГруппы расходятся:")
        print(f"  эталон: {gg}")
        print(f"  карта:  {cg}")
    if missing:
        print(f"\nНет в карте ({len(missing)}): {', '.join(sorted(missing, key=_key))}")
    if extra:
        print(f"\nЛишние в карте ({len(extra)}): {', '.join(sorted(extra, key=_key))}")
    if diff:
        print(f"\nРазошлись значения ({len(diff)}):")
        for no, g, c in sorted(diff, key=lambda t: _key(t[0])):
            print(f"  {no:>6}: эталон «{g}» / карта «{c}»")
    return len(missing) + len(extra) + len(diff)


def _key(no: str) -> tuple[int, int]:
    m = re.match(r"(\d+)(?:-(\d+))?", no)
    return (int(m.group(1)), int(m.group(2) or 0)) if m else (9999, 0)


def main() -> None:
    ap = argparse.ArgumentParser(description="Печать и сверка карт обмера")
    ap.add_argument("--card", type=Path, required=True)
    ap.add_argument("--gold", type=Path)
    ap.add_argument("--quiet", action="store_true", help="только сверка, без печати таблицы")
    args = ap.parse_args()

    title, rows = read_card(args.card)
    if not args.quiet:
        dump(title, rows)

    if args.gold:
        _, gold_rows = read_card(args.gold)
        problems = compare(rows, gold_rows)
        sys.exit(1 if problems else 0)


if __name__ == "__main__":
    main()
