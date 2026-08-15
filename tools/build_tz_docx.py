"""Сборка требований по изготовлению (.docx) из tz.json.

Оформление берётся из `templates/tz.docx` — шаблон вырезан из реальной работы,
поэтому шрифты, поля, маркированные и нумерованные списки совпадают с привычными.
Каждый блок технических требований получает собственный номерной список,
иначе нумерация в Word пойдёт сквозной через весь документ.

    python tools/build_tz_docx.py work/tz.json -o "output/... (ТЗ).docx"
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu

TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "tz.docx"

ABSTRACT_BULLET = "4"      # маркированный список — «Общие данные»
ABSTRACT_DECIMAL = "0"     # нумерованный — технические требования

TITLE = "ТРЕБОВАНИЯ ПО ИЗГОТОВЛЕНИЮ"
PARTS_DOC_TITLE = "ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ИЗГОТОВЛЕНИЕ"
PARTS_TITLE = "ТРЕБОВАНИЯ К ДЕТАЛЯМ (ПОЛНАЯ ДЕТАЛИЗАЦИЯ)"
SPEC_HEADER = ["№ на сборке", "Обозначение", "Наименование", "Кол-во", "Примечание"]
SPEC_WIDTHS = [781685, 1612265, 2104390, 635635, 804545]
SKIP_SECTION_HEADER = "Документация"

# Таблица параметров зубчатого венца (ГОСТ 2.403): наименование, обозначение, значение.
PARAM_WIDTHS = [3200000, 900000, 1840000]

# Тонкая светлая сетка — как в эталонных заданиях.
TABLE_PROPERTIES = (
    f'<w:tblPr {nsdecls("w")}>'
    '<w:tblW w:w="0" w:type="auto"/>'
    '<w:tblCellSpacing w:w="15" w:type="dxa"/>'
    '<w:tblBorders>'
    '<w:top w:val="single" w:sz="2" w:space="0" w:color="E3E3E3"/>'
    '<w:left w:val="single" w:sz="2" w:space="0" w:color="E3E3E3"/>'
    '<w:bottom w:val="single" w:sz="2" w:space="0" w:color="E3E3E3"/>'
    '<w:right w:val="single" w:sz="2" w:space="0" w:color="E3E3E3"/>'
    '<w:insideH w:val="single" w:sz="2" w:space="0" w:color="E3E3E3"/>'
    '<w:insideV w:val="single" w:sz="2" w:space="0" w:color="E3E3E3"/>'
    '</w:tblBorders>'
    '<w:tblCellMar><w:left w:w="0" w:type="dxa"/><w:right w:w="0" w:type="dxa"/></w:tblCellMar>'
    '<w:tblLook w:val="04A0" w:firstRow="1" w:lastRow="0" w:firstColumn="1" '
    'w:lastColumn="0" w:noHBand="0" w:noVBand="1"/>'
    '</w:tblPr>'
)


def _new_list(doc, abstract_id: str) -> str:
    """Новый экземпляр списка: нумерация каждого блока начинается с единицы."""
    numbering = doc.part.numbering_part.element
    used = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = str(max(used) + 1 if used else 1)

    element = numbering.makeelement(qn("w:num"), {qn("w:numId"): num_id})
    link = numbering.makeelement(qn("w:abstractNumId"), {qn("w:val"): abstract_id})
    element.append(link)
    numbering.append(element)
    return num_id


def _para(doc, text: str, bold: bool = False, num_id: str | None = None):
    paragraph = doc.add_paragraph()
    if num_id is not None:
        properties = paragraph._p.get_or_add_pPr()
        num_pr = properties.makeelement(qn("w:numPr"), {})
        level = properties.makeelement(qn("w:ilvl"), {qn("w:val"): "0"})
        link = properties.makeelement(qn("w:numId"), {qn("w:val"): num_id})
        num_pr.append(level)
        num_pr.append(link)
        properties.append(num_pr)
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def _tech(doc, items: list[dict]) -> None:
    if not items:
        return
    num_id = _new_list(doc, ABSTRACT_DECIMAL)
    for item in items:
        _para(doc, item["text"] if isinstance(item, dict) else str(item),
              num_id=num_id)


def _spec_table(doc, rows: list[dict]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(SPEC_HEADER))
    current = table._tbl.find(qn("w:tblPr"))
    if current is not None:
        table._tbl.replace(current, parse_xml(TABLE_PROPERTIES))

    for cell, text, width in zip(table.rows[0].cells, SPEC_HEADER, SPEC_WIDTHS):
        cell.width = Emu(width)
        cell.paragraphs[0].add_run(text).bold = True

    section = ""
    for row in rows:
        if row["section"] != section:
            section = row["section"]
            if section and section != SKIP_SECTION_HEADER:
                cells = table.add_row().cells
                cells[2].paragraphs[0].add_run(section).bold = True
        cells = table.add_row().cells
        for cell, text, width in zip(
                cells,
                [row["pos"], row["designation"], row["name"], row["qty"], row["note"]],
                SPEC_WIDTHS):
            cell.width = Emu(width)
            if text:
                cell.paragraphs[0].add_run(text)


def _full_name(node: dict) -> str:
    return " ".join(x for x in (node.get("designation"), node.get("title")) if x)


def _table(doc, widths: list[int], rows: list[list[str]], header: bool = False) -> None:
    table = doc.add_table(rows=0, cols=len(widths))
    current = table._tbl.find(qn("w:tblPr"))
    if current is not None:
        table._tbl.replace(current, parse_xml(TABLE_PROPERTIES))
    for number, row in enumerate(rows):
        cells = table.add_row().cells
        for cell, text, width in zip(cells, row, widths):
            cell.width = Emu(width)
            if text:
                cell.paragraphs[0].add_run(text).bold = header and number == 0


def _parameters(doc, parameters: list[dict]) -> None:
    if not parameters:
        return
    _table(doc, PARAM_WIDTHS,
           [[p.get("name", ""), p.get("symbol", ""), p.get("value", "")]
            for p in parameters])


def _part_block(doc, part: dict) -> None:
    _para(doc, _full_name(part), bold=True)
    _tech(doc, part["tech_requirements"])
    facts = []
    if part.get("mass"):
        facts.append(f"Вес – {part['mass']}")
    if part.get("material"):
        facts.append(f"Материал – {part['material']}")
    if facts:
        _para(doc, " | ".join(facts), bold=True)


def build_parts(tz: dict, out: Path, template: Path = TEMPLATE) -> int:
    """Задание на комплект деталей: сборки нет, у каждой детали свой чертёж."""
    doc = docx.Document(str(template))
    _para(doc, PARTS_DOC_TITLE, bold=True)
    _para(doc, tz["title"], bold=True)

    if tz.get("general"):
        _para(doc, "Общие данные:", bold=True)
        num_id = _new_list(doc, ABSTRACT_BULLET)
        for line in tz["general"]:
            _para(doc, line, num_id=num_id)

    for part in tz["parts"]:
        head = _full_name(part)
        if part.get("qty"):
            head += f" — {part['qty']}"
        _para(doc, head, bold=True)
        if part.get("material"):
            _para(doc, f"Материал – {part['material']}", bold=True)
        _parameters(doc, part.get("parameters", []))
        if part.get("tech_requirements"):
            _para(doc, "Технические требования:", bold=True)
            _tech(doc, part["tech_requirements"])

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return len(doc.paragraphs)


def _assembly_block(doc, node: dict, number: int) -> None:
    _para(doc, f"{number}. {_full_name(node)}", bold=True)
    if node["tech_requirements"]:
        _para(doc, "Технические требования:", bold=True)
        _tech(doc, node["tech_requirements"])
    if node["mass"]:
        _para(doc, f"Вес – {node['mass']}", bold=True)
    if node["spec"]:
        _para(doc, "Спецификация:", bold=True)
        _spec_table(doc, node["spec"])
    for part in node["parts"]:
        _part_block(doc, part)


def build(tz: dict, out: Path, template: Path = TEMPLATE) -> int:
    if tz.get("kind") == "parts":
        return build_parts(tz, out, template)

    doc = docx.Document(str(template))
    root = tz["assembly"]
    header = _full_name(tz)

    _para(doc, TITLE, bold=True)
    _para(doc, header, bold=True)

    if tz.get("general"):
        _para(doc, "Общие данные:", bold=True)
        num_id = _new_list(doc, ABSTRACT_BULLET)
        for line in tz["general"]:
            _para(doc, line, num_id=num_id)

    _para(doc, "Технические требования чертежей", bold=True)
    _para(doc, header, bold=True)
    _tech(doc, root["tech_requirements"])
    if root["mass"]:
        _para(doc, f"Вес – {root['mass']}", bold=True)

    _para(doc, "Спецификация", bold=True)
    _para(doc, header, bold=True)
    _spec_table(doc, root["spec"])

    for number, child in enumerate(root["children"], 1):
        _assembly_block(doc, child, number)

    if root["parts"]:
        _para(doc, PARTS_TITLE, bold=True)
        _para(doc, "Детали основной сборки:", bold=True)
        for part in root["parts"]:
            _part_block(doc, part)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return len(doc.paragraphs)


def main() -> None:
    ap = argparse.ArgumentParser(description="Требования по изготовлению в Word из tz.json")
    ap.add_argument("tz", type=Path)
    ap.add_argument("-o", "--out", type=Path, required=True)
    ap.add_argument("--template", type=Path, default=TEMPLATE)
    args = ap.parse_args()

    tz = json.loads(args.tz.read_text(encoding="utf-8"))
    count = build(tz, args.out, args.template)
    print(f"абзацев: {count} -> {args.out}")


if __name__ == "__main__":
    main()
